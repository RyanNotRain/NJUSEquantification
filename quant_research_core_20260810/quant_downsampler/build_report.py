"""Build the final project report from the current, leakage-safe outputs."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT.parent / "output"
REPORT = ROOT / "兆信基金量化研究完整结果报告_最终版.docx"
ASSETS = OUT / "report_assets"

# standard_business_brief preset, with Microsoft YaHei as the East Asian font.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
RISK = "9B1C1C"
GOLD = "7A5A00"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def _set_cell_width(cell, width):
    properties = cell._tc.get_or_add_tcPr()
    width_tag = properties.find(qn("w:tcW"))
    if width_tag is None:
        width_tag = OxmlElement("w:tcW")
        properties.append(width_tag)
    width_tag.set(qn("w:w"), str(width))
    width_tag.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        tag = margins.find(qn(f"w:{name}"))
        if tag is None:
            tag = OxmlElement(f"w:{name}")
            margins.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def _shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    properties.append(tag)


def _set_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    _repeat_header(header)
    for cell, text in zip(header.cells, headers):
        _shade(cell, LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(str(text)), size=font_size, bold=True, color=DARK_BLUE)
    for row_values in rows:
        cells = table.add_row().cells
        for column, (cell, value) in enumerate(zip(cells, row_values)):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.05
            set_font(paragraph.add_run(str(value)), size=font_size)
    _set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_paragraph(doc, text="", bold=False, color=None, size=11, italic=False,
                  align=None, before=0, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if align is not None:
        paragraph.alignment = align
    set_font(paragraph.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return paragraph


def add_bullet(doc, text, color=None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    set_font(paragraph.add_run(text), size=11, color=color)
    return paragraph


def add_callout(doc, label, text, color=DARK_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE_BLUE)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    properties.append(borders)
    set_font(paragraph.add_run(f"{label}："), bold=True, color=color)
    set_font(paragraph.add_run(text), color=INK)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_font(run, size=16 if level == 1 else 13 if level == 2 else 12,
                 bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return paragraph


def add_picture(doc, path, caption, alt_text, width=6.15):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    paragraph.paragraph_format.keep_with_next = True
    cap = add_paragraph(
        doc, caption, size=9.2, color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8,
    )
    cap.paragraph_format.keep_together = True


def add_page_number(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(paragraph.add_run("第 "), size=9, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    set_font(paragraph.add_run(" 页"), size=9, color=GRAY)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, before, after, color in (
        (1, 16, 16, 8, BLUE), (2, 13, 12, 6, BLUE), (3, 12, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167


def build_factor_stability_chart(stability):
    from PIL import Image, ImageDraw, ImageFont

    ASSETS.mkdir(parents=True, exist_ok=True)
    pivot = stability.pivot(index="factor", columns="split", values="IC")
    order = pivot.get("train", pd.Series(index=pivot.index, dtype=float)).sort_values().index
    pivot = pivot.reindex(order)
    width, height = 1700, 1020
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 24)
    small = ImageFont.truetype(str(font_path), 20)
    title_font = ImageFont.truetype(str(bold_path), 32)
    label_font = ImageFont.truetype(str(font_path), 21)
    left, right, top, bottom = 390, 1620, 115, 930
    x_min, x_max = -0.12, 0.12

    def x_position(value):
        return left + (float(value) - x_min) / (x_max - x_min) * (right - left)

    draw.text((width // 2, 35), "Factor IC stability by chronological split",
              fill="#202124", font=title_font, anchor="ma")
    for tick in (-0.12, -0.08, -0.04, 0.00, 0.04, 0.08, 0.12):
        x = x_position(tick)
        draw.line((x, top, x, bottom), fill="#DADCE0", width=2)
        draw.text((x, bottom + 18), f"{tick:.2f}", fill="#555555", font=small, anchor="ma")
    zero_x = x_position(0)
    draw.line((zero_x, top, zero_x, bottom), fill="#202124", width=3)
    colors = {"train": "#2E74B5", "validation": "#8FAADC", "test": "#C55A11"}
    row_height = (bottom - top) / len(pivot)
    bar_height = 7
    offsets = {"train": -12, "validation": 0, "test": 12}
    for row_number, (factor, values) in enumerate(pivot.iterrows()):
        center_y = top + row_height * (row_number + 0.5)
        draw.text((left - 18, center_y), factor, fill="#202124", font=label_font, anchor="rm")
        for split in ("train", "validation", "test"):
            value = values.get(split)
            if pd.isna(value):
                continue
            value_x = x_position(value)
            y = center_y + offsets[split]
            draw.rectangle(
                (min(zero_x, value_x), y - bar_height, max(zero_x, value_x), y + bar_height),
                fill=colors[split],
            )
    legend_x = left
    for split in ("train", "validation", "test"):
        draw.rectangle((legend_x, 78, legend_x + 24, 96), fill=colors[split])
        draw.text((legend_x + 34, 87), split, fill="#202124", font=small, anchor="lm")
        legend_x += 180
    draw.text(((left + right) // 2, height - 28), "Mean IC", fill="#202124", font=font, anchor="ma")
    path = ASSETS / "factor_stability_ic.png"
    image.save(path, quality=95)
    return path


def pct(value):
    return f"{float(value):.2%}"


def main():
    factor_summary = pd.read_csv(OUT / "evaluation" / "factor_summary.csv")
    stability = pd.read_csv(OUT / "evaluation" / "factor_stability.csv")
    backtest = pd.read_csv(OUT / "backtest_strict" / "metrics.csv").set_index("strategy")
    static_models = json.loads((OUT / "factor_models_aligned" / "results.json").read_text(encoding="utf-8"))
    rolling_models = json.loads((OUT / "factor_models_rolling_aligned" / "results.json").read_text(encoding="utf-8"))
    classifier = json.loads((OUT / "lstm_next_minute" / "test_metrics.json").read_text(encoding="utf-8"))
    classifier_config = json.loads((OUT / "lstm_next_minute" / "config.json").read_text(encoding="utf-8"))
    hierarchical = json.loads(
        (OUT / "lstm_next_minute_hierarchical" / "test_metrics.json").read_text(encoding="utf-8")
    )
    hierarchical_config = json.loads(
        (OUT / "lstm_next_minute_hierarchical" / "config.json").read_text(encoding="utf-8")
    )
    stability_chart = build_factor_stability_chart(stability)
    adaptive = backtest.loc["adaptive_close"]
    rolling_recent_return = rolling_models["models"]["return_regression"]["last_45_days"]
    rolling_recent_sharpe = rolling_models["models"]["direct_sharpe"]["last_45_days"]

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("兆信基金量化研究 | 最终结果报告"), size=9, color=GRAY)
    add_page_number(section)

    # Editorial cover header pattern.
    add_paragraph(doc, "量化研究项目", size=12, bold=True, color=BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=80, after=18)
    add_paragraph(doc, "兆信基金笔试题", size=29, bold=True, color=INK,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_paragraph(doc, "完整结果报告（最终修订版）", size=17, color=DARK_BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    add_paragraph(
        doc, "逐笔降采样 · 因子评价 · 无前视回测 · 因子学习 · LSTM",
        size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=70,
    )
    add_table(doc, ["样本区间", "股票范围", "报告日期"], [[
        "2025-04-01 至 2026-06-30", "300 只匿名股票", str(date.today()),
    ]], [3360, 2880, 3120], font_size=10)
    add_paragraph(
        doc, "研究说明：adjfactor.pkl 的六位代码可与逐笔文件完整一一映射。"
             "OHLC 已按归一化复权因子调整，成交量、成交额与成交笔数保留实际交易口径。",
        size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=0,
    )
    doc.add_page_break()

    add_heading(doc, "一、执行摘要")
    add_callout(
        doc, "可信主结论",
        f"项目已按复权、242 分钟口径完整跑通。严格 Task4 中 adaptive_close 年化复合收益 "
        f"{pct(adaptive['annual_return_cagr'])}、夏普 {adaptive['sharpe_ratio']:.3f}、"
        f"最大回撤 {pct(adaptive['max_drawdown'])}；Task5 已把平盘纳入训练，三分类全分钟准确率 "
        f"{pct(classifier['accuracy'])}、平衡准确率 {pct(classifier['balanced_accuracy'])}、"
        f"Macro F1 {pct(classifier['macro_f1'])}。",
    )
    add_table(doc, ["模块", "状态", "最重要结果"], [
        ["Task 1", "通过", "302 个交易日、300 股票；复权后分钟表 242×300"],
        ["Task 2–3", "完成", "10 因子；全样本 IC 全为负，测试期多数发生符号反转"],
        ["Task 4", "完成", f"严格 adaptive_close 夏普 {adaptive['sharpe_ratio']:.3f}；Naive 仅作泄漏对照"],
        ["因子学习", "完成", f"最近 45 日两模型夏普 {rolling_recent_return['annualized_sharpe_net']:.2f} / {rolling_recent_sharpe['annualized_sharpe_net']:.2f}"],
        ["Task 5", "完成", f"下跌/平盘/上涨三分类；全分钟 balanced accuracy {pct(classifier['balanced_accuracy'])}"],
    ], [1500, 1300, 6560])
    add_bullet(doc, "负 IC 本身不是错误：可对因子取反；稳定性和可获得时点才是决定能否交易的关键。")
    add_bullet(doc, "题目字面策略使用未来收益计算当前 IC，必然泄露。报告保留 Naive 结果只为量化偏差，不把它当作策略成绩。")
    add_bullet(doc, "复权因子可由六位代码映射到带交易所后缀的证券代码；早先的“无法映射”判断已纠正。")

    add_heading(doc, "二、数据处理与 Task 1")
    add_paragraph(
        doc, "原始逐笔记录包含 Time、Price、Volume、BSFlag。价格除以 100 还原；"
        "开盘集合竞价参与日频 OHLC，但不进入连续竞价分钟表。分钟频显示 09:30–11:30 "
        "与 13:00–15:00，共 242 行；14:57–14:59 为收盘集合竞价等待行，15:00 保留撮合结果。",
    )
    add_table(doc, ["频率", "输出形状", "指标", "填充规则"], [
        ["日频", "302×300", "复权 OHLC、量、笔数、额、主买/主卖量额", "OHLC 前值；成交类 0"],
        ["分钟频", "每指标 302 张 242×300", "同上，共 11 类", "连续竞价 OHLC 前值；成交类 0"],
    ], [1200, 2100, 3100, 2960])
    add_callout(
        doc, "验收结果", "冒烟测试全部通过；300 个逐笔代码与复权因子 300 个证券代码完整重合；"
        "随机抽样日期的分钟结果与队友完整版本逐值一致。",
    )

    add_heading(doc, "三、Task 2–3：因子构建、方向与稳定性")
    add_paragraph(
        doc, "因子包括题目示例因子、5/10/20 日动量、5/10/20 日波动率、买卖失衡、"
        "日内振幅和 20 日量价相关。评价收益只保留信号日和下一交易日均有成交的股票，"
        "避免停牌价格前向填充产生伪零收益。",
    )
    factor_rows = []
    for _, row in factor_summary.iterrows():
        factor_rows.append([
            row["factor"], f"{row['IC']:.4f}", f"{row['IR']:.3f}",
            f"{row['rank_IC']:.4f}", str(int(row["n_days"])),
        ])
    add_table(doc, ["因子", "IC", "IR/ICIR", "Rank IC", "有效日"], factor_rows,
              [3400, 1300, 1500, 1600, 1560], font_size=8.6)
    add_picture(
        doc, stability_chart, "图 1  因子 IC 的训练/验证/测试时间稳定性",
        "十个因子在按时间划分的训练、验证和测试区间中的平均 IC 对比",
    )
    add_callout(
        doc, "稳定性风险",
        "训练期和验证期大多为负；测试期除买卖失衡外，多数因子转为正。"
        "这解释了为什么用早期历史训练的反向权重在最后阶段失效。",
        color=RISK,
    )

    add_heading(doc, "四、Task 4：严格 IC 加权回测")
    add_heading(doc, "4.1 时点修正", level=2)
    add_paragraph(
        doc, "信号在交易日 t 收盘后形成，于 t+1 的开盘或收盘执行。用于评价因子 t 的收益是"
        "执行价 t+1 到 t+2 的收益；因此该 IC 最早在 t+2 才完整可知。严格策略将原始 IC "
        "至少延迟两个交易日，再做 rolling、expanding、hybrid 或 adaptive 加权。",
    )
    add_table(doc, ["步骤", "已知信息", "处理"], [
        ["t 日收盘", "当日因子", "生成信号，不使用尚未发生的收益"],
        ["t+1", "执行价格、可交易状态", "按开盘或收盘建仓；停牌股票不能交易"],
        ["t+2", "完整持有期收益", "因子 t 的 IC 此时才可进入后续权重"],
    ], [1500, 3400, 4460])
    add_heading(doc, "4.2 组合与成本", level=2)
    add_paragraph(
        doc, "初始资金 1,000 万元，每日选择信号最高的 10 只股票。持仓和现金逐日记账；"
        "停牌持仓锁定，不允许按前向填充价格卖出；卖出手续费万分之五，买入手续费不计。"
        "流动性过滤只使用决策时点已知的成交量和过去 20 日平均成交额。",
    )
    strict_names = [
        "rolling_close", "expanding_close", "hybrid_close", "adaptive_close",
        "rolling_open", "expanding_open", "hybrid_open", "adaptive_open",
    ]
    strict_rows = []
    for name in strict_names:
        row = backtest.loc[name]
        strict_rows.append([
            name, pct(row["annual_return_cagr"]), f"{row['sharpe_ratio']:.3f}",
            pct(row["max_drawdown"]), pct(row["average_sell_turnover"]),
        ])
    add_table(doc, ["策略", "年化复合", "夏普", "最大回撤", "平均卖出换手"], strict_rows,
              [2700, 1700, 1300, 1700, 1960], font_size=8.7)
    add_picture(
        doc, OUT / "backtest_strict" / "nav_curve.png",
        "图 2  无前视偏差的 8 组 Task4 策略净值",
        "仅包含可交易历史 IC 版本的八条净值曲线，不含 Naive 泄露对照",
    )
    naive_rows = []
    for name in ("naive_close", "naive_open"):
        row = backtest.loc[name]
        naive_rows.append([
            name, pct(row["annual_return_cagr"]), f"{row['sharpe_ratio']:.2f}",
            f"{row['final_nav']:.2f}", pct(row["max_drawdown"]),
        ])
    add_table(doc, ["泄露对照", "年化复合", "夏普", "期末净值", "最大回撤"], naive_rows,
              [2200, 1700, 1300, 2200, 1960])
    add_picture(
        doc, OUT / "backtest_strict" / "lookahead_bias_comparison.png",
        "图 3  Naive 与严格策略对照（纵轴为对数尺度）",
        "严格自适应策略与使用未来收益的 Naive 策略净值差异",
    )
    add_callout(
        doc, "为什么旧图像空白",
        "Naive 净值超过 20 倍，而严格策略约 1.1 倍；放在同一线性纵轴会把严格曲线压扁。"
        "新版把可信策略单独绘图，泄露对照使用对数坐标。",
        color=GOLD,
    )

    add_heading(doc, "五、收益率训练与直接夏普训练")
    add_paragraph(
        doc, "两套模型都只使用 10 个日频因子。收益率模型最小化个股未来收益误差；"
        "直接夏普模型学习因子权重，最大化扣除 5 bp 换手成本后的组合夏普。"
        "收益目标统一为因子 t 后，在 t+1 开盘建仓、t+2 开盘退出。",
    )
    static_rows = []
    for model in ("return_regression", "direct_sharpe"):
        for split in ("train", "validation", "test"):
            value = static_models["models"][model][split]
            static_rows.append([
                model, split, f"{value['annualized_sharpe_net']:.2f}",
                pct(value["annualized_return_net"]), pct(value["max_drawdown"]),
            ])
    add_table(doc, ["模型", "区间", "净夏普", "年化净收益", "最大回撤"], static_rows,
              [2600, 1600, 1500, 1900, 1760], font_size=8.8)
    rolling_rows = []
    for model in ("return_regression", "direct_sharpe"):
        values = rolling_models["models"][model]
        rolling_rows.append([
            model, f"{values['full_walk_forward']['annualized_sharpe_net']:.2f}",
            pct(values["full_walk_forward"]["annualized_return_net"]),
            f"{values['last_45_days']['annualized_sharpe_net']:.2f}",
            pct(values["last_45_days"]["annualized_return_net"]),
        ])
    add_table(doc, ["滚动模型", "全期夏普", "全期年化", "最近45日夏普", "最近45日年化"], rolling_rows,
              [2600, 1600, 1700, 1800, 1660])
    add_picture(
        doc, OUT / "factor_models_rolling_aligned" / "portfolio_nav.png",
        "图 4  120 日滚动训练的两套因子模型净值（已扣假设成本）",
        "收益率回归和直接夏普优化在滚动样本外区间的净值曲线",
    )
    add_callout(
        doc, "模型判断",
        f"全滚动期看起来很好，但最近 45 日两者夏普分别为 "
        f"{rolling_recent_return['annualized_sharpe_net']:.2f} 和 {rolling_recent_sharpe['annualized_sharpe_net']:.2f}。"
        "必须把近期反转作为主风险，不能只汇报全期指标。",
        color=RISK,
    )

    add_heading(doc, "六、Task 5：下一分钟下跌/平盘/上涨 LSTM 分类")
    add_heading(doc, "6.1 方法与严格时间切分", level=2)
    add_paragraph(
        doc, "主实验沿用队友版本中有效的设定：对 5 只股票使用截至当前时刻的 60 分钟序列，"
        "预测下一分钟收盘价相对当前分钟的方向。输入为 11 个平稳化特征，包括开高低相对收盘价、"
        "收盘对数收益，以及成交量、成交笔数、成交额和主买/主卖订单流的 log1p 变换。"
        "上午和下午分别构造序列，不跨越午休；收盘集合竞价等待段不用于建模。",
    )
    add_paragraph(
        doc, "训练集截至 2026-05-31，验证集为 2026-06-01 至 06-15，测试集为 06-16 至 06-30。"
        "标准化均值和方差只由训练集拟合；平盘作为独立类别参与训练和测试。采用平方根逆频率权重，"
        "它比完全逆频率更能识别平盘，又比不加权版本保持更好的三类平衡表现。主评价覆盖全部有效分钟。",
    )
    add_table(doc, ["股票", "训练/验证/测试样本", "Accuracy", "Balanced Acc.", "Macro F1"], [[
        str(len(classifier_config["stock_codes"])),
        f"{classifier_config['sizes']['train']} / {classifier_config['sizes']['validation']} / {classifier_config['sizes']['test']}",
        pct(classifier["accuracy"]), pct(classifier["balanced_accuracy"]), pct(classifier["macro_f1"]),
    ]], [1100, 3000, 1700, 1900, 1660])
    add_picture(
        doc, OUT / "lstm_next_minute" / "training_history.png",
        "图 5  下一分钟 LSTM 训练损失与验证指标",
        "下一分钟三分类模型十二轮训练的损失、验证准确率和验证平衡准确率",
    )
    add_heading(doc, "6.2 三类表现与交易诊断", level=2)
    class_rows = []
    for name, label in (("down", "下跌"), ("flat", "平盘"), ("up", "上涨")):
        values = classifier["per_class"][name]
        class_rows.append([
            label, str(values["support"]), pct(values["precision"]),
            pct(values["recall"]), pct(values["f1"]),
        ])
    add_table(doc, ["类别", "样本数", "Precision", "Recall", "F1"], class_rows,
              [1800, 1500, 2000, 2000, 2060], font_size=9.0)
    add_table(doc, ["诊断指标", "结果", "解释"], [
        ["平盘多数基线", pct(classifier["majority_baseline"]), "全预测为样本最多的平盘类"],
        ["真实非平盘方向准确率", pct(classifier["nonflat_direction_accuracy"]), "预测平盘在该口径计为错误"],
        ["模型预测涨跌比例", pct(classifier["predicted_move_rate"]), "模型在全部分钟中发出方向信号的比例"],
        ["涨跌信号精确命中率", pct(classifier["move_signal_exact_accuracy"]), "预测涨跌时必须与真实三分类完全一致"],
    ], [3000, 1700, 4660], font_size=8.8)
    add_callout(
        doc, "Task5 结论",
        f"三分类全分钟准确率 {pct(classifier['accuracy'])}，高于平盘多数基线 "
        f"{pct(classifier['majority_baseline'])}；平衡准确率 {pct(classifier['balanced_accuracy'])} 和 "
        f"Macro F1 {pct(classifier['macro_f1'])} 表明模型对三类都学到了一定信息。"
        f"但平盘召回率仅 {pct(classifier['per_class']['flat']['recall'])}，涨跌信号精确命中率仅 "
        f"{pct(classifier['move_signal_exact_accuracy'])}，仍不足以直接支持扣成本交易。",
    )
    add_heading(doc, "6.3 两阶段优化实验", level=2)
    add_paragraph(
        doc, "进一步测试了共享编码器的两阶段模型：移动头使用全部样本判断平盘/变动，"
        "方向头只在真实非平盘样本上学习下跌/上涨。输入额外加入 5/15/60 分钟平盘比例、"
        "连续未变价时长、短期绝对收益、实现波动率和分钟位置，并只在验证集校准两个阈值。",
    )
    add_table(doc, ["模型", "特征数", "Accuracy", "Balanced Acc.", "Macro F1", "移动检测 Bal."], [
        ["直接三分类（正式）", str(len(classifier_config["feature_names"])), pct(classifier["accuracy"]),
         pct(classifier["balanced_accuracy"]), pct(classifier["macro_f1"]),
         pct(classifier["movement_detection_balanced_accuracy"])],
        ["两阶段共享 LSTM", str(len(hierarchical_config["feature_names"])), pct(hierarchical["accuracy"]),
         pct(hierarchical["balanced_accuracy"]), pct(hierarchical["macro_f1"]),
         pct(hierarchical["movement_detection_balanced_accuracy"])],
    ], [2150, 1000, 1450, 1710, 1450, 1600], font_size=8.4)
    add_callout(
        doc, "优化判断",
        f"两阶段模型把移动检测平衡准确率从 {pct(classifier['movement_detection_balanced_accuracy'])} "
        f"提高到 {pct(hierarchical['movement_detection_balanced_accuracy'])}，但最终 Macro F1 从 "
        f"{pct(classifier['macro_f1'])} 降到 {pct(hierarchical['macro_f1'])}。"
        "说明拆任务方向正确，但共享编码器存在任务冲突；因此保留实验输出，不替换正式模型。",
        color=GOLD,
    )

    add_heading(doc, "七、限制与下一步")
    add_bullet(doc, "复权：代码映射已解决，但本项目只对 OHLC 复权；成交量和成交额仍按真实交易量额解释。")
    add_bullet(doc, "交易：回测未模拟涨跌停、冲击成本、成交容量和盘口成交概率。")
    add_bullet(doc, "因子：最近阶段符号反转明显；应加入短期 regime 检测、行业/规模中性化和因子正交化。")
    add_bullet(doc, "验证：应延长样本并采用多个滚动测试窗，避免一次 45 日测试偶然性。")
    add_bullet(doc, "LSTM：共享两阶段实验改善了移动检测但未改善最终三分类；下一步可尝试两个独立编码器或门控专家模型。")

    add_heading(doc, "八、复现命令")
    add_table(doc, ["命令", "用途"], [
        ["run_all.bat", "Task 1–4、严格回测、两套因子模型"],
        ["run_all.bat --with-lstm", "在上述流程后训练并评估 Task5 LSTM"],
        ["py -3 scripts\\smoke_test.py", "快速验收降采样逻辑"],
        ["py -3 scripts\\run_task4_strict.py", "只重新生成严格 Task4 结果和图"],
    ], [4300, 5060], font_size=9.2)
    add_paragraph(
        doc, "正式数据与结果目录：project/output。旧回测、旧 LSTM 回归和重复调参结果已清理。",
        size=9.5, color=GRAY, after=0,
    )

    doc.core_properties.title = "兆信基金量化研究完整结果报告（最终修订版）"
    doc.core_properties.subject = "Task 1–5 全链路方法、结果与风险"
    doc.core_properties.author = "量化研究项目"
    doc.core_properties.keywords = "量化研究, 因子, IC, 回测, LSTM"
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
