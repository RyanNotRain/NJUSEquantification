"""Build the corrected, Task4-only detailed report."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from build_report import (
    BLUE, DARK_BLUE, GRAY, INK, RISK, add_bullet, add_callout, add_heading,
    add_page_number, add_paragraph, add_picture, add_table, configure_styles,
    pct, set_font,
)


ROOT = Path(__file__).resolve().parent
OUT = ROOT.parent / "output"
REPORT = ROOT / "Task4_回测方法与结果详细报告.docx"


def main():
    metrics = pd.read_csv(OUT / "backtest_strict" / "metrics.csv").set_index("strategy")
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("兆信基金 | Task 4 严格回测报告"), size=9, color=GRAY)
    add_page_number(section)

    add_paragraph(doc, "TASK 4 / RESEARCH NOTE", size=11, bold=True, color=BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=42, after=16)
    add_paragraph(doc, "因子 IC 加权回测", size=26, bold=True, color=INK,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_paragraph(doc, "时点修正、真实持仓记账与前视偏差对照", size=15, color=DARK_BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_table(doc, ["初始资金", "持仓数", "卖出手续费", "执行方式"], [[
        "1,000 万元", "10", "万分之五", "次日开盘 / 次日收盘",
    ]], [2100, 1600, 2200, 3460])
    add_callout(
        doc, "结果摘要", f"严格版本最优夏普为 adaptive_close：年化复合收益 "
        f"{pct(metrics.loc['adaptive_close', 'annual_return_cagr'])}，夏普 "
        f"{metrics.loc['adaptive_close', 'sharpe_ratio']:.3f}，最大回撤 "
        f"{pct(metrics.loc['adaptive_close', 'max_drawdown'])}。Naive 版本年化超过 1200%，"
        "但它使用未来收益确定当前权重，不能视为策略能力。",
    )

    add_heading(doc, "一、题目策略的问题")
    add_paragraph(
        doc, "题目要求使用“当日因子值与次日收益率计算的 IC”给当日因子加权。"
        "在当日形成信号时，次日收益尚未发生，因此这个 IC 不可获得。若直接拿它选股，"
        "等价于先看答案再交易，产生前视偏差。",
    )
    add_bullet(doc, "当日因子可以在当日收盘后得到；对应收益至少要等后续持有期结束。")
    add_bullet(doc, "不能把全样本平均 IC 直接用于整个样本，因为它同样包含未来区间。")
    add_bullet(doc, "Naive 结果应保留为错误对照，但必须与可信结果分开绘图和解释。")

    add_heading(doc, "二、严格时点链路")
    add_table(doc, ["时点", "发生事项", "策略允许使用的信息"], [
        ["t 日收盘", "计算当日 10 个因子", "只用截至 t 的数据形成信号"],
        ["t+1 开盘/收盘", "执行买卖", "结合当日成交量判断是否可交易"],
        ["t+2 开盘/收盘", "持有期收益完整实现", "因子 t 的 IC 此时才进入历史权重"],
    ], [1900, 3300, 4160])
    add_paragraph(
        doc, "实现中先把原始逐日 IC 延迟两个交易日，再构造 rolling、expanding、hybrid "
        "和 adaptive 权重。adaptive 同时要求近期与长期 IC 方向一致、近期绝对 IC 不低于 0.01，"
        "并用历史 IC 波动率调整强度。",
    )

    add_heading(doc, "三、组合与交易约束")
    add_bullet(doc, "每日按组合信号选择前 10 只股票，等资金目标权重。")
    add_bullet(doc, "现金和每只股票股数逐日保存；不是把每日选股收益简单平均。")
    add_bullet(doc, "执行日成交量为 0 的股票不能买卖；已持有的停牌股票继续锁定。")
    add_bullet(doc, "决策日使用过去 20 日平均成交额过滤流动性，不读取未来成交额。")
    add_bullet(doc, "卖出手续费万分之五，买入手续费不计；收益指标按 252 日年化。")

    add_heading(doc, "四、严格回测结果")
    names = [
        "rolling_close", "expanding_close", "hybrid_close", "adaptive_close",
        "rolling_open", "expanding_open", "hybrid_open", "adaptive_open",
    ]
    rows = []
    for name in names:
        row = metrics.loc[name]
        rows.append([
            name, pct(row["annual_return_cagr"]), pct(row["annual_volatility"]),
            f"{row['sharpe_ratio']:.3f}", pct(row["max_drawdown"]),
            pct(row["average_sell_turnover"]),
        ])
    add_table(doc, ["策略", "年化复合", "年化波动", "夏普", "最大回撤", "卖出换手"], rows,
              [2460, 1400, 1400, 1100, 1500, 1500], font_size=8.5)
    add_picture(
        doc, OUT / "backtest_strict" / "nav_curve.png",
        "图 1  八组无前视偏差策略净值",
        "滚动、扩展、混合和自适应历史 IC 在开盘和收盘执行下的净值曲线",
    )
    add_callout(
        doc, "表现解释", "收盘执行整体好于开盘执行，自适应过滤改善了夏普和回撤。"
        "不过 2026 年 5 月后所有版本明显回落，与 Task3 中因子方向反转一致。",
    )

    add_heading(doc, "五、Naive 对照与旧图空白原因")
    rows = []
    for name in ("naive_close", "naive_open"):
        row = metrics.loc[name]
        rows.append([
            name, pct(row["annual_return_cagr"]), f"{row['sharpe_ratio']:.2f}",
            f"{row['final_nav']:.2f}", pct(row["max_drawdown"]),
        ])
    add_table(doc, ["对照", "年化复合", "夏普", "期末净值", "最大回撤"], rows,
              [2200, 1700, 1300, 2200, 1960])
    add_picture(
        doc, OUT / "backtest_strict" / "lookahead_bias_comparison.png",
        "图 2  Naive 与自适应策略对照（对数坐标）",
        "使用未来收益的 Naive 策略与严格自适应策略的数量级差异",
    )
    add_paragraph(
        doc, "旧图把 20 倍以上的 Naive 净值和约 1.1 倍的严格净值放在同一线性纵轴，"
        "严格曲线被压到坐标轴附近，看起来像没有线。现在主图只画严格策略，"
        "偏差对照图改用对数纵轴。",
    )

    add_heading(doc, "六、结论与限制")
    add_callout(
        doc, "最终判断", "Task4 的方法链路已修正，严格回测不再是全面亏损；"
        "但夏普低于 1、回撤约 16%–24%，且近期明显反转，只能算有一定研究信号，"
        "不能宣称策略已经成熟。",
        color=RISK,
    )
    add_bullet(doc, "六位代码已与 adjfactor.pkl 完整映射，OHLC 使用归一化复权因子；量额保持真实交易口径。")
    add_bullet(doc, "尚未加入涨跌停、冲击成本、容量和盘口成交概率。")
    add_bullet(doc, "未来应以多个滚动样本外窗口为主，而不是继续针对单一全期指标调参。")
    add_paragraph(
        doc, "结果来源：project/output/backtest_strict/metrics.csv、日度净值、IC 权重、"
        "选股明细及两张净值图。",
        size=9.5, color=GRAY, before=8, after=0,
    )

    doc.core_properties.title = "Task 4 因子 IC 加权回测详细报告"
    doc.core_properties.author = "量化研究项目"
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
